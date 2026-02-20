#!/usr/bin/env python3
"""
Check generated resume files (PDF/DOCX) for ATS compliance.
"""

import argparse
import re
import subprocess
import sys
from pathlib import Path

from docx import Document


EMOJI_RE = re.compile(
    "["
    "\U0001f600-\U0001f64f"
    "\U0001f300-\U0001f5ff"
    "\U0001f680-\U0001f6ff"
    "\U0001f1e0-\U0001f1ff"
    "\U00002702-\U000027b0"
    "\U0001f900-\U0001f9ff"
    "\U00002600-\U000026ff"
    "\U0000fe00-\U0000fe0f"
    "\U0000200d"
    "\U00002060"
    "\U0000231a-\U0000231b"
    "]+"
)

ATS_SAFE_FONTS = {"Calibri", "Arial", "Times New Roman", "Helvetica", "Georgia", "Cambria"}


def check_docx(path):
    """Run ATS compliance checks on a DOCX file."""
    doc = Document(path)
    results = []

    # 1. Heading styles
    builtin_headings = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            builtin_headings.append((p.style.name, p.text[:60]))
    if builtin_headings:
        results.append(("PASS", "Heading styles", f"{len(builtin_headings)} built-in headings used"))
    else:
        results.append(("WARN", "Heading styles", "No built-in heading styles found - ATS may not parse sections"))

    # 2. Fonts
    fonts_used = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
    default_font = doc.styles["Normal"].font.name
    unsafe = fonts_used - ATS_SAFE_FONTS
    if unsafe:
        results.append(("WARN", "Fonts", f"Non-standard fonts: {', '.join(unsafe)}"))
    else:
        label = f"{default_font}" + (f" + {', '.join(fonts_used - {default_font})}" if fonts_used - {default_font} else "")
        results.append(("PASS", "Fonts", f"ATS-safe: {label}"))

    # 3. Tables
    if doc.tables:
        results.append(("WARN", "Tables", f"{len(doc.tables)} table(s) found - some ATS cannot parse these"))
    else:
        results.append(("PASS", "Tables", "None"))

    # 4. Images
    image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    if image_count:
        results.append(("WARN", "Images", f"{image_count} image(s) found - ATS cannot read image text"))
    else:
        results.append(("PASS", "Images", "None"))

    # 5. Text colors
    colors = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                colors.add(str(run.font.color.rgb))
    non_standard = {c for c in colors if c not in ("000000", "0000FF")}
    if non_standard:
        results.append(("WARN", "Text colors", f"Non-black/non-link colors: {', '.join(non_standard)}"))
    else:
        results.append(("PASS", "Text colors", "Black text only" + (" + hyperlink blue" if "0000FF" in colors else "")))

    # 6. Emoji
    emoji_found = any(EMOJI_RE.search(p.text) for p in doc.paragraphs)
    if emoji_found:
        results.append(("WARN", "Emoji", "Emoji characters found - some ATS strip or misparse these"))
    else:
        results.append(("PASS", "Emoji", "None"))

    # 7. Layout (columns)
    multi_col = False
    for section in doc.sections:
        cols = section._sectPr.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols")
        if cols:
            num = cols[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num", "1")
            if int(num) > 1:
                multi_col = True
    if multi_col:
        results.append(("WARN", "Layout", "Multi-column layout detected - use single column for ATS"))
    else:
        results.append(("PASS", "Layout", "Single column"))

    # 8. Hyperlinks
    link_count = sum(1 for rel in doc.part.rels.values() if "hyperlink" in rel.reltype)
    if link_count:
        results.append(("PASS", "Hyperlinks", f"{link_count} preserved as real links"))
    else:
        results.append(("INFO", "Hyperlinks", "None found"))

    # 9. Style usage summary
    style_counts = {}
    for p in doc.paragraphs:
        style_counts[p.style.name] = style_counts.get(p.style.name, 0) + 1
    styles_summary = ", ".join(f"{k}: {v}" for k, v in sorted(style_counts.items()))
    results.append(("INFO", "Styles used", styles_summary))

    return results


def check_pdf(path):
    """Run ATS compliance checks on a PDF file."""
    results = []

    # 1. File size
    size_kb = path.stat().st_size / 1024
    if size_kb > 2048:
        results.append(("WARN", "File size", f"{size_kb:.0f} KB - may exceed some ATS upload limits"))
    else:
        results.append(("PASS", "File size", f"{size_kb:.0f} KB"))

    # 2. Text extraction (requires pdftotext)
    pdf_text = None
    try:
        proc = subprocess.run(
            ["pdftotext", str(path), "-"],
            capture_output=True, text=True, timeout=10,
        )
        if proc.returncode == 0 and proc.stdout.strip():
            pdf_text = proc.stdout
            lines = [l for l in pdf_text.strip().split("\n") if l.strip()]
            results.append(("PASS", "Text extraction", f"{len(lines)} lines extracted - text is selectable"))
        else:
            results.append(("WARN", "Text extraction", "pdftotext returned no text - may be image-based"))
    except FileNotFoundError:
        results.append(("SKIP", "Text extraction", "pdftotext not installed (apt install poppler-utils)"))
    except subprocess.TimeoutExpired:
        results.append(("WARN", "Text extraction", "pdftotext timed out"))

    # 3. Emoji in extracted text
    if pdf_text:
        if EMOJI_RE.search(pdf_text):
            results.append(("WARN", "Emoji", "Emoji characters found in PDF text"))
        else:
            results.append(("PASS", "Emoji", "None"))

    # 4. Standard sections present
    if pdf_text:
        expected = ["experience", "education", "skills"]
        found = [s for s in expected if s in pdf_text.lower()]
        missing = [s for s in expected if s not in pdf_text.lower()]
        if missing:
            results.append(("WARN", "Sections", f"Missing standard sections: {', '.join(missing)}"))
        else:
            results.append(("PASS", "Sections", f"Found: {', '.join(found)}"))

    # 5. Page count (via pdfinfo)
    try:
        proc = subprocess.run(
            ["pdfinfo", str(path)],
            capture_output=True, text=True, timeout=10,
        )
        if proc.returncode == 0:
            for line in proc.stdout.split("\n"):
                if line.startswith("Pages:"):
                    pages = int(line.split(":")[1].strip())
                    if pages > 2:
                        results.append(("WARN", "Page count", f"{pages} pages - consider condensing to 1-2"))
                    else:
                        results.append(("PASS", "Page count", f"{pages}"))
                    break
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass  # pdfinfo not available, skip silently

    return results


def print_results(file_path, results):
    """Print check results in a formatted table."""
    print(f"\n{'=' * 60}")
    print(f"  ATS Compliance: {file_path}")
    print(f"{'=' * 60}")

    pass_count = sum(1 for r in results if r[0] == "PASS")
    warn_count = sum(1 for r in results if r[0] == "WARN")
    check_count = sum(1 for r in results if r[0] in ("PASS", "WARN"))

    for status, check, detail in results:
        icon = {"PASS": "[PASS]", "WARN": "[WARN]", "SKIP": "[SKIP]", "INFO": "[INFO]"}[status]
        print(f"  {icon}  {check}: {detail}")

    print(f"{'─' * 60}")
    if warn_count == 0:
        print(f"  Result: {pass_count}/{check_count} checks passed - ATS ready")
    else:
        print(f"  Result: {pass_count}/{check_count} checks passed, {warn_count} warning(s)")
    print()


def main():
    parser = argparse.ArgumentParser(description="Check resume files for ATS compliance")
    parser.add_argument("files", nargs="+", help="PDF or DOCX files to check")
    args = parser.parse_args()

    exit_code = 0
    for file_arg in args.files:
        path = Path(file_arg)
        if not path.exists():
            print(f"Error: File '{path}' not found")
            exit_code = 1
            continue

        if path.suffix.lower() == ".docx":
            results = check_docx(path)
        elif path.suffix.lower() == ".pdf":
            results = check_pdf(path)
        else:
            print(f"Error: Unsupported file type '{path.suffix}' (use .pdf or .docx)")
            exit_code = 1
            continue

        print_results(path, results)
        if any(r[0] == "WARN" for r in results):
            exit_code = 1

    return exit_code


if __name__ == "__main__":
    sys.exit(main())
