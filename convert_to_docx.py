#!/usr/bin/env python3
"""
Convert Markdown resume to ATS-friendly DOCX format.
Uses Word built-in styles (Heading 1, Heading 2, etc.) for ATS parser compatibility.
"""

import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from bs4 import BeautifulSoup


# Regex to strip emoji characters
EMOJI_RE = re.compile(
    "["
    "\U0001f600-\U0001f64f"
    "\U0001f300-\U0001f5ff"
    "\U0001f680-\U0001f6ff"
    "\U0001f1e0-\U0001f1ff"
    "\U00002702-\U000027b0"
    "\U0001f900-\U0001f9ff"
    "\U00002600-\U000026ff"
    "\U0000fe00-\U0000fe0f"
    "\U0000200d"
    "\U00002060"
    "\U0000231a-\U0000231b"
    "]+",
    flags=re.UNICODE,
)


def strip_emojis(text):
    """Remove emoji characters from text."""
    return EMOJI_RE.sub("", text).strip()


def parse_markdown(md_file):
    """Read and parse markdown file."""
    with open(md_file, "r", encoding="utf-8") as f:
        content = f.read()

    html = markdown2.markdown(content, extras=["fenced-code-blocks", "tables"])
    return BeautifulSoup(html, "html.parser")


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = paragraph._element.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink",
        attrib={"{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id": r_id},
    )

    new_run = paragraph.add_run(text)
    new_run.font.color.rgb = RGBColor(0, 0, 255)
    new_run.font.underline = True

    hyperlink.append(new_run._element)
    paragraph._element.append(hyperlink)

    return hyperlink


def process_text_with_links(paragraph, element):
    """Process text that may contain links."""
    if element.name == "a":
        add_hyperlink(paragraph, element.get_text(), element.get("href"))
    else:
        for child in element.children:
            if isinstance(child, str):
                if child.strip():
                    paragraph.add_run(child)
            elif child.name == "a":
                add_hyperlink(paragraph, child.get_text(), child.get("href"))
            else:
                process_text_with_links(paragraph, child)


def create_docx(soup, output_file):
    """Create ATS-friendly DOCX document from parsed HTML."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for element in soup.children:
        if isinstance(element, str):
            continue

        if element.name == "h1":
            # Use built-in Heading 1 style for ATS
            text = strip_emojis(element.get_text())
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif element.name == "p":
            text = strip_emojis(element.get_text())
            if "|" in text and "@" in text:  # Contact info line
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                parts = text.split("|")
                for i, part in enumerate(parts):
                    part = part.strip()

                    # Check if this part contains a link
                    link_element = None
                    for child in element.children:
                        if child.name == "a" and child.get_text().strip() in part:
                            link_element = child
                            break

                    if link_element:
                        before_link = part.split(link_element.get_text())[0]
                        if before_link:
                            p.add_run(before_link)
                        add_hyperlink(p, link_element.get_text(), link_element.get("href"))
                        after_parts = part.split(link_element.get_text())
                        after_link = after_parts[1] if len(after_parts) > 1 else ""
                        if after_link:
                            p.add_run(after_link)
                    else:
                        p.add_run(part)

                    if i < len(parts) - 1:
                        p.add_run(" | ")

                if p.runs:
                    p.runs[0].font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                process_text_with_links(p, element)

        elif element.name == "h2":
            # Use built-in Heading 2 style for ATS - black text
            text = strip_emojis(element.get_text())
            p = doc.add_heading(text, level=2)
            p.space_after = Pt(6)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif element.name == "h3":
            # Use built-in Heading 3 style for ATS
            text = strip_emojis(element.get_text())
            p = doc.add_heading(text, level=3)
            p.space_after = Pt(3)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif element.name == "em":
            p = doc.add_paragraph()
            run = p.add_run(strip_emojis(element.get_text()))
            run.font.italic = True
            run.font.size = Pt(10)

        elif element.name == "ul":
            for li in element.find_all("li"):
                p = doc.add_paragraph(style="List Bullet")

                strong_tags = li.find_all("strong")
                if strong_tags:
                    text = str(li)
                    text = text.replace("<li>", "").replace("</li>", "")

                    parts = re.split(r"(<strong>.*?</strong>)", text)
                    for part in parts:
                        if part.startswith("<strong>"):
                            bold_text = part.replace("<strong>", "").replace("</strong>", "")
                            run = p.add_run(strip_emojis(bold_text))
                            run.font.bold = True
                        else:
                            p.add_run(strip_emojis(part))
                else:
                    p.add_run(strip_emojis(li.get_text()))

                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(3)

    doc.save(output_file)


def main():
    parser = argparse.ArgumentParser(description="Convert markdown resume to ATS-friendly DOCX")
    parser.add_argument("markdown_file", help="Path to markdown resume file")
    parser.add_argument(
        "-o",
        "--output-dir",
        help="Output directory (default: outputs/ in project directory)",
        default=None,
    )

    args = parser.parse_args()

    md_file = Path(args.markdown_file)
    if not md_file.exists():
        print(f"Error: File '{md_file}' not found")
        sys.exit(1)

    if args.output_dir:
        output_dir = Path(args.output_dir)
    else:
        script_dir = Path(__file__).parent
        output_dir = script_dir / "outputs"

    output_dir.mkdir(parents=True, exist_ok=True)

    soup = parse_markdown(md_file)
    output_file = output_dir / md_file.with_suffix(".docx").name
    create_docx(soup, output_file)

    print(f"Successfully created: {output_file}")


if __name__ == "__main__":
    main()
